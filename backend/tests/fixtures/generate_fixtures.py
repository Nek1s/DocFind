"""Генератор набора тестовых документов для проверки загрузки и парсинга (QA-03).

Скрипт детерминированно собирает фикстуры в подпапки рядом с собой:

* ``valid/``    — корректные PDF/DOCX разных размеров и структур;
* ``complex/``  — сложные, но валидные: таблицы, изображения, нестандартные
  шрифты, «поехавшее» форматирование;
* ``negative/`` — отбраковываемые валидацией: пустые, с подменённым расширением,
  zip без ``word/document.xml``;
* ``corrupt/``  — синтаксически битые PDF/DOCX, которые ломают извлечение текста.

Бинарные фикстуры закоммичены в репозиторий (ТЗ разрешает бинарники только в
``tests/fixtures``). Генератор нужен для воспроизводимого пересоздания набора и
заодно документирует, как именно собран каждый файл. Файл ``sample.pdf``
(каноничный многостраничный PDF с текстовым слоем) создан ранее и здесь не
перегенерируется.

Маркеры (известный текст) и списки путей экспортируются как единый источник
правды для проверочного теста ``tests/test_fixtures.py``.

Зависимости генерации: ``python-docx`` (есть в requirements.txt) и ``Pillow``
(см. requirements-dev.txt).

Запуск::

    python generate_fixtures.py [--with-oversized]
"""
import argparse
import io
import zipfile
from pathlib import Path

import docx
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw

FIXTURES_DIR = Path(__file__).parent

# --- Известный текст («маркеры»), который тесты ищут в извлечённом содержимом ---
SIMPLE_TEXT = "Простой корректный документ DocFind для проверки извлечения текста."
TABLE_HEADING = "Экзаменационная ведомость"
TABLE_HEADER = ("Студент", "Группа", "Оценка")
TABLE_ROW = ("Иванов И.И.", "ПИ-101", "5")
IMAGE_CAPTION = "Документ с встроенным изображением для проверки парсера."
FONT_TEXT = "Текст нестандартным шрифтом — извлекаться должен как обычный."
FONT_NAME = "OCR A Extended"  # заведомо нестандартный шрифт
LARGE_MARKER = "Большой документ DocFind"
BROKEN_MARKER = "Документ со сломанным форматированием"
SCAN_CAPTION = "СКАН страница"  # рисуется на картинке, в текстовый слой НЕ попадает

# Размер, который в проекте считается предельным для загрузки (BE-02). Дублируем
# константу осознанно: генератор не должен тянуть зависимость от кода приложения.
MAX_UPLOAD_SIZE = 20 * 1024 * 1024  # 20 МБ

# --- Каталог путей по категориям (относительно FIXTURES_DIR) — для тестов ---
VALID_DOCS = ("sample.pdf", "valid/simple.docx", "valid/large.docx")
# Валидные, но без извлекаемого текстового слоя (скан) — отдельная проверка.
VALID_NO_TEXT = ("valid/scanned.pdf",)
COMPLEX_DOCS = (
    "complex/with_table.docx",
    "complex/with_image.docx",
    "complex/nonstandard_font.docx",
    "complex/broken_formatting.docx",
)
# Должны отбраковываться `validate_document` (формат/размер/содержимое).
NEGATIVE_DOCS = (
    "negative/empty.pdf",
    "negative/empty.docx",
    "negative/pdf_renamed_to.docx",
    "negative/docx_renamed_to.pdf",
    "negative/text_renamed_to.pdf",
    "negative/archive_without_word.docx",
)
# Битые: проходят (или нет) валидацию, но ломают извлечение текста.
CORRUPT_DOCS = ("corrupt/truncated.pdf", "corrupt/not_a_zip.docx")
OVERSIZED_DOC = "negative/oversized.pdf"  # >20 МБ, не коммитится (см. .gitignore)


def _png_bytes(text: str, size: tuple[int, int] = (320, 120)) -> bytes:
    """Маленький PNG с подписью — для встраивания в DOCX."""
    image = Image.new("RGB", size, "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((4, 4, size[0] - 5, size[1] - 5), outline="black", width=2)
    draw.text((20, 50), text, fill="black")
    buf = io.BytesIO()
    image.save(buf, format="PNG")
    return buf.getvalue()


def _save_docx(document: docx.document.Document, rel_path: str) -> None:
    """Сохранить документ python-docx в фикстуру по относительному пути."""
    out = FIXTURES_DIR / rel_path
    out.parent.mkdir(parents=True, exist_ok=True)
    document.save(out)


def _write_bytes(rel_path: str, content: bytes) -> None:
    """Записать произвольные байты в фикстуру по относительному пути."""
    out = FIXTURES_DIR / rel_path
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_bytes(content)


# --- Корректные документы ---------------------------------------------------


def build_simple_docx() -> None:
    """Короткий валидный DOCX из заголовка и пары абзацев."""
    doc = docx.Document()
    doc.add_heading("DocFind — тестовый документ", level=1)
    doc.add_paragraph(SIMPLE_TEXT)
    doc.add_paragraph("Второй абзац с дополнительным текстом для чанкирования.")
    _save_docx(doc, "valid/simple.docx")


def build_large_docx() -> None:
    """Объёмный валидный DOCX (~сотни КБ) — проверка работы на больших файлах."""
    doc = docx.Document()
    doc.add_heading(LARGE_MARKER, level=1)
    for i in range(3000):
        doc.add_paragraph(
            f"Абзац {i}. Лекция по информатике: алгоритмы, структуры данных, "
            f"сложность вычислений и практические примеры применения."
        )
    _save_docx(doc, "valid/large.docx")


def build_scanned_pdf() -> None:
    """PDF из картинок без текстового слоя — имитация скана (текст не извлечётся)."""
    pages = []
    for n in (1, 2):
        page = Image.new("RGB", (1240, 1754), "white")  # ~A4 при 150 dpi
        draw = ImageDraw.Draw(page)
        draw.text((120, 140), f"{SCAN_CAPTION} {n}", fill="black")
        pages.append(page)
    out = FIXTURES_DIR / "valid/scanned.pdf"
    out.parent.mkdir(parents=True, exist_ok=True)
    pages[0].save(out, "PDF", save_all=True, append_images=pages[1:], resolution=150)


# --- Сложные, но валидные документы -----------------------------------------


def build_table_docx() -> None:
    """DOCX с таблицей. NB: текущий парсер (paragraphs) текст таблиц НЕ извлекает."""
    doc = docx.Document()
    doc.add_heading(TABLE_HEADING, level=1)
    table = doc.add_table(rows=1, cols=3)
    for cell, name in zip(table.rows[0].cells, TABLE_HEADER):
        cell.text = name
    for cell, value in zip(table.add_row().cells, TABLE_ROW):
        cell.text = value
    _save_docx(doc, "complex/with_table.docx")


def build_image_docx() -> None:
    """DOCX с встроенным растровым изображением и текстовой подписью."""
    doc = docx.Document()
    doc.add_heading("Документ с изображением", level=1)
    doc.add_paragraph(IMAGE_CAPTION)
    doc.add_picture(io.BytesIO(_png_bytes("DocFind")), width=Inches(2))
    _save_docx(doc, "complex/with_image.docx")


def build_nonstandard_font_docx() -> None:
    """DOCX, где текст набран нестандартным шрифтом (на извлечение не влияет)."""
    doc = docx.Document()
    doc.add_heading("Нестандартный шрифт", level=1)
    run = doc.add_paragraph().add_run(FONT_TEXT)
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.font.bold = True
    _save_docx(doc, "complex/nonstandard_font.docx")


def build_broken_formatting_docx() -> None:
    """Валидный DOCX с «поехавшим» форматированием: пустые абзацы, вложенная таблица."""
    doc = docx.Document()
    doc.add_heading(BROKEN_MARKER, level=1)
    for _ in range(15):
        doc.add_paragraph("")  # пачка пустых абзацев
    doc.add_paragraph("   \t  смешанные    пробелы   и\tтабуляции   ")
    outer = doc.add_table(rows=1, cols=1)
    inner = outer.rows[0].cells[0].add_table(rows=1, cols=2)  # таблица в таблице
    inner.rows[0].cells[0].text = "вложенная"
    inner.rows[0].cells[1].text = "таблица"
    doc.add_paragraph(f"{BROKEN_MARKER} — конец")
    _save_docx(doc, "complex/broken_formatting.docx")


# --- Негативные кейсы (должны отбраковываться валидацией) --------------------


def build_negative_docs() -> None:
    """Пустые файлы, подменённые расширения и zip без word/document.xml."""
    _write_bytes("negative/empty.pdf", b"")
    _write_bytes("negative/empty.docx", b"")

    # Подменённое расширение: реальный PDF под именем .docx и наоборот.
    pdf_bytes = (FIXTURES_DIR / "sample.pdf").read_bytes()
    _write_bytes("negative/pdf_renamed_to.docx", pdf_bytes)
    docx_bytes = (FIXTURES_DIR / "valid/simple.docx").read_bytes()
    _write_bytes("negative/docx_renamed_to.pdf", docx_bytes)
    _write_bytes("negative/text_renamed_to.pdf", "Просто текст, не PDF.".encode())

    # Валидный zip, но без word/document.xml (например .xlsx под видом .docx).
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as archive:
        archive.writestr("xl/workbook.xml", "<workbook/>")
    _write_bytes("negative/archive_without_word.docx", buf.getvalue())


# --- Битые документы (ломают извлечение текста) -----------------------------


def build_corrupt_docs() -> None:
    """PDF/DOCX с верной сигнатурой, но битой структурой."""
    # Сигнатура %PDF на месте (валидация пройдёт), но структура PDF разрушена —
    # извлечение текста должно упасть TextExtractionError.
    _write_bytes("corrupt/truncated.pdf", b"%PDF-1.4\n\xde\xad\xbe\xef not a real pdf body")
    # Сигнатура zip (PK) есть, но это не валидный архив — не DOCX.
    _write_bytes(
        "corrupt/not_a_zip.docx",
        b"PK\x03\x04\xff\xfe " + "это не настоящий docx".encode(),
    )


def build_oversized_pdf() -> None:
    """PDF чуть больше лимита загрузки (>20 МБ). Не коммитится — см. .gitignore."""
    _write_bytes(OVERSIZED_DOC, b"%PDF-1.4\n" + b"\x00" * (MAX_UPLOAD_SIZE + 1024))


def main() -> None:
    """Собрать весь набор фикстур. С ``--with-oversized`` создаёт и файл >20 МБ."""
    parser = argparse.ArgumentParser(description="Генерация тестовых документов (QA-03).")
    parser.add_argument(
        "--with-oversized",
        action="store_true",
        help="дополнительно создать negative/oversized.pdf (>20 МБ, не коммитится)",
    )
    args = parser.parse_args()

    # Порядок важен: негативные кейсы переиспользуют уже собранные valid-файлы.
    build_simple_docx()
    build_large_docx()
    build_scanned_pdf()
    build_table_docx()
    build_image_docx()
    build_nonstandard_font_docx()
    build_broken_formatting_docx()
    build_negative_docs()
    build_corrupt_docs()
    if args.with_oversized:
        build_oversized_pdf()

    print(f"Готово: фикстуры собраны в {FIXTURES_DIR}")


if __name__ == "__main__":
    main()
